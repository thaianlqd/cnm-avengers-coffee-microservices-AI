from docx import Document
from docx.shared import Pt

OUTPUT = "UseCase_KhachVangLai.docx"

DOC_PATH = OUTPUT

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level==1 else 12)

def add_kv(doc, key, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{key}\t")
    run.bold = True
    p.add_run(value)

def add_basic_flow_table(doc, steps):
    # steps: list of (actor_step, system_step)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Khách vãng lai').bold = True
    hdr_cells[1].paragraphs[0].add_run('Hệ thống').bold = True
    for a,s in steps:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(a)
        row[1].paragraphs[0].add_run(s)
    doc.add_paragraph()


def add_listed_section(doc, title, desc, precond, postcond, actor_main, actor_sub, basic_steps, alternatives, exceptions, subflows, entity_map):
    add_heading(doc, title, level=1)
    add_kv(doc, 'Mô tả', desc)
    add_kv(doc, 'Tiền điều kiện', '\n'.join(['- '+x for x in precond]))
    add_kv(doc, 'Hậu điều kiện', '\n'.join(['- '+x for x in postcond]))
    add_kv(doc, 'Actor chính', actor_main)
    add_kv(doc, 'Actor phụ', actor_sub)
    doc.add_paragraph('Basic flow')
    add_basic_flow_table(doc, basic_steps)
    if alternatives:
        doc.add_paragraph('Alternative flow')
        for k,v in alternatives.items():
            doc.add_paragraph(f"{k}:")
            for line in v:
                doc.add_paragraph(line, style='List Number')
    if subflows:
        doc.add_paragraph('Subflows')
        for k,v in subflows.items():
            doc.add_paragraph(f"{k}:")
            for line in v:
                doc.add_paragraph(line, style='List Number')
    if exceptions:
        doc.add_paragraph('Exception')
        for k,v in exceptions.items():
            doc.add_paragraph(f"{k}:")
            for line in v:
                doc.add_paragraph(line, style='List Number')
    # Entity mapping
    doc.add_paragraph('Mapping tới Domain (đọc / tạo / cập nhật / xóa)')
    for mode, items in entity_map.items():
        doc.add_paragraph(f"{mode}:")
        for ent in items:
            doc.add_paragraph(f"- {ent}")
    doc.add_page_break()


def main():
    doc = Document()
    doc.core_properties.title = 'Đặc tả Use Case - Khách vãng lai'

    # 1. Xem danh sách sản phẩm / Chi tiết sản phẩm
    add_listed_section(
        doc,
        '1. Xem danh sách sản phẩm / Chi tiết sản phẩm',
        'Cho phép khách vãng lai duyệt danh sách sản phẩm theo danh mục và xem thông tin chi tiết từng sản phẩm. Thông tin hiển thị bao gồm tên, giá, hình ảnh, mô tả, nhãn nổi bật/mới và trạng thái còn hàng theo từng chi nhánh.',
        ['Hệ thống đang hoạt động bình thường.', 'Dữ liệu sản phẩm và danh mục đã có trong CSDL, ít nhất một sản phẩm đang được bán.', 'Khách vãng lai đã truy cập trang chủ ứng dụng.'],
        ['Danh sách hoặc trang chi tiết được hiển thị; không có ghi thay đổi vào CSDL; nếu không có sản phẩm hiển thị thông báo phù hợp.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Truy cập trang danh sách sản phẩm qua menu hoặc banner trang chủ.', 'Lấy danh sách sản phẩm đang bán từ CSDL cùng tên danh mục tương ứng.'),
            ('Kiểm tra danh sách có dữ liệu hay không.', 'Trả về kết quả (có/không).'),
            ('Hiển thị danh sách sản phẩm dạng lưới (hình ảnh, tên, giá, danh mục, nhãn).', ''),
            ('Chọn tab danh mục để lọc nhanh.', 'Lấy lại danh sách sản phẩm theo danh mục đã chọn.'),
            ('Chọn một sản phẩm để xem chi tiết.', 'Kiểm tra sản phẩm tồn tại và đang bán; lấy thêm thông tin tồn kho theo chi nhánh; trả về dữ liệu chi tiết.'),
            ('Xem trang chi tiết và có thể nhấn "Quay lại" để trở về danh sách.', '')
        ],
        {
            '3.1 Không có sản phẩm nào đang bán': ['Hệ thống hiển thị thông báo "Hiện chưa có sản phẩm nào.".', 'Kết thúc use case.'],
            '7.1 Danh mục được chọn chưa có sản phẩm': ['Hệ thống hiển thị thông báo "Danh mục này chưa có sản phẩm.".', 'Quay lại bước 4.'],
            '9.1 Sản phẩm không tồn tại hoặc đã ngừng bán': ['Hệ thống hiển thị thông báo "Sản phẩm không còn tồn tại hoặc đã ngừng kinh doanh.".', 'Quay lại bước 1.']
        },
        {
            '2.1 Lỗi kết nối CSDL khi tải danh sách': ['Hiển thị thông báo lỗi và ghi log; kết thúc use case.'],
            '9.1 Lỗi CSDL khi tải chi tiết sản phẩm': ['Hiển thị thông báo lỗi và ghi log; kết thúc use case.']
        },
        {'None': []},
        {
            'Đọc': ['san_pham (Sản phẩm)', 'danh_muc (Danh mục)', 'ton_kho_san_pham (Tồn kho theo chi nhánh)'],
            'Tạo': [],
            'Cập nhật': [],
            'Xóa': []
        }
    )

    # 2. Tìm kiếm & Lọc sản phẩm
    add_listed_section(
        doc,
        '2. Tìm kiếm & Lọc sản phẩm',
        'Cho phép khách vãng lai tìm kiếm sản phẩm theo từ khóa (tên) và áp dụng bộ lọc đa tiêu chí: danh mục, khoảng giá, nhãn nổi bật, sản phẩm mới hoặc đang có khuyến mãi. Kết quả hiển thị theo đúng tiêu chí đã chọn.',
        ['Hệ thống đang hoạt động bình thường.', 'Dữ liệu sản phẩm và danh mục đã có trong CSDL.', 'Khách vãng lai đang ở màn hình danh sách sản phẩm.'],
        ['Kết quả tìm kiếm / lọc được hiển thị đúng; không có ghi thay đổi vào CSDL; bộ lọc được duy trì khi điều hướng.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Nhập từ khóa và/hoặc chọn bộ lọc (danh mục, khoảng giá, nhãn).', 'Kiểm tra đầu vào; chuẩn hóa từ khóa; tìm kiếm sản phẩm trong CSDL theo điều kiện; chỉ lấy sản phẩm đang bán.'),
            ('Kiểm tra kết quả có dữ liệu hay không.', 'Trả về danh sách kết quả và tổng số.'),
            ('Điều chỉnh bộ lọc; hệ thống cập nhật lại kết quả.', ''),
            ('Chọn sản phẩm để xem chi tiết.', 'Chuyển sang use case Xem chi tiết sản phẩm.')
        ],
        {
            '2.1 Không nhập từ khóa và không chọn bộ lọc nào': ['Hệ thống hiển thị toàn bộ danh sách sản phẩm đang bán.', 'Quay lại bước 6.'],
            '2.2 Khoảng giá không hợp lệ': ['Hiển thị thông báo "Khoảng giá không hợp lệ, vui lòng nhập lại."', 'Quay lại bước 1.'],
            '5.1 Không có sản phẩm nào khớp': ['Hiển thị thông báo "Không tìm thấy sản phẩm phù hợp."', 'Đề xuất thử từ khóa khác hoặc xóa bộ lọc; quay lại bước 1.']
        },
        {
            '4.1 Lỗi CSDL khi thực hiện tìm kiếm': ['Hiển thị thông báo lỗi, ghi log, kết thúc use case.']
        },
        {'None': []},
        {
            'Đọc': ['san_pham', 'danh_muc', 'khuyen_mai (để hiển thị nhãn khuyến mãi nếu cần)'],
            'Tạo': [],
            'Cập nhật': [],
            'Xóa': []
        }
    )

    # 3. Đăng ký tài khoản
    add_listed_section(
        doc,
        '3. Đăng ký tài khoản',
        'Cho phép khách vãng lai tạo tài khoản mới bằng email/số điện thoại hoặc qua Google để sử dụng tính năng đặt hàng và tích điểm. Hệ thống tự động tạo hồ sơ khách hàng ban đầu.',
        ['Hệ thống đang hoạt động bình thường.', 'Email và số điện thoại chưa được đăng ký trong CSDL.'],
        ['Tài khoản và hồ sơ khách hàng được tạo; email xác nhận được gửi; nếu thất bại thì không tạo.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Chọn chức năng "Đăng ký" trên giao diện.', 'Hiển thị form đăng ký hoặc nút đăng ký bằng Google.'),
            ('Nhập thông tin (Họ tên, SĐT, Email, Mật khẩu) hoặc chọn Google.', 'Kiểm tra định dạng các trường; nếu chọn Google, chuyển luồng xác thực Google.'),
            ('Nhấn "Đăng ký".', 'Kiểm tra email/số điện thoại chưa tồn tại; mã hóa mật khẩu; lưu `nguoi_dung` vào CSDL với vai trò Khách hàng và trạng thái Đang hoạt động.'),
            ('Hệ thống tự động tạo `ho_so_khach_hang` (điểm=0, hạng=Thành viên).', 'Gửi email xác nhận; hiển thị thông báo "Đăng ký thành công" và chuyển sang màn hình đăng nhập.')
        ],
        {
            '4.1 Thiếu trường bắt buộc': ['Hiển thị thông báo yêu cầu nhập đủ thông tin.', 'Quay lại bước 3.'],
            '6.1 Email đã được đăng ký': ['Hiển thị "Email này đã được sử dụng."', 'Quay lại bước 3.'],
            '7.1 SĐT đã được đăng ký': ['Hiển thị "Số điện thoại này đã được sử dụng."', 'Quay lại bước 3.']
        },
        {
            '5.1 Khách hủy đăng ký': ['Hiển thị xác nhận; nếu xác nhận hủy thì kết thúc use case, không lưu.'],
            '9.1 Lỗi lưu vào CSDL': ['Hiển thị lỗi; huỷ thao tác; ghi log; kết thúc.'],
            '11.1 Lỗi gửi email': ['Tài khoản vẫn được tạo; hiển thị thông báo lỗi gửi email; ghi log.']
        },
        {
            'Subflow: Đăng ký bằng Google': [
                'Khách nhấn "Đăng ký bằng Google".',
                'Hệ thống chuyển sang xác thực Google; nếu Google trả về email/họ tên, kiểm tra email đã tồn tại hay chưa.',
                'Nếu chưa tồn tại, tiếp tục luồng tạo tài khoản như trên; nếu đã tồn tại, thông báo và quay lại bước nhập.'
            ]
        },
        {
            'Đọc': [],
            'Tạo': ['nguoi_dung', 'ho_so_khach_hang'],
            'Cập nhật': [],
            'Xóa': []
        }
    )

    # small note: fix avatar display caveat
    doc.add_paragraph('Ghi chú: Khi tạo tài khoản mới, nếu email/SĐT khớp tài khoản khác, không tự động hiển thị avatar cũ của tài khoản đó (tạm giữ avatar mặc định).')

    # 4. Xem trang tin tức, bài viết
    add_listed_section(
        doc,
        '4. Xem trang tin tức, bài viết',
        'Cho phép khách vãng lai xem danh sách bài viết đã xuất bản và đọc nội dung chi tiết từng bài. Mỗi lần xem chi tiết, lượt xem tăng +1.',
        ['Hệ thống hoạt động.', 'Dữ liệu bài viết đã có trong CSDL với trạng thái xuất bản.'],
        ['Danh sách hoặc bài viết chi tiết được hiển thị; lượt xem bài viết được cập nhật +1 trong CSDL.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Truy cập mục Tin tức trên thanh điều hướng.', 'Lấy danh sách bài viết đã xuất bản từ CSDL, sắp xếp theo ngày đăng.'),
            ('Kiểm tra danh sách có dữ liệu không.', ''),
            ('Hiển thị danh sách (hình ảnh, tiêu đề, sapo, chuyên mục, tác giả, ngày, lượt xem).', ''),
            ('Chọn bài viết để đọc chi tiết.', 'Kiểm tra bài viết tồn tại và trạng thái xuất bản; tăng lượt xem +1; trả về nội dung chi tiết.')
        ],
        {
            '3.1 Không có bài viết đã xuất bản': ['Hiển thị "Hiện chưa có bài viết nào."', 'Kết thúc use case.'],
            '7.1 Không có bài viết trong chuyên mục': ['Hiển thị thông báo tương ứng', 'Quay lại bước 5.']
        },
        {
            '2.1 Lỗi CSDL khi tải danh sách': ['Hiển thị lỗi, ghi log, kết thúc.'],
            '11.1 Lỗi cập nhật lượt xem': ['Bỏ qua lỗi cập nhật lượt xem, ghi log, tiếp tục hiển thị nội dung.']
        },
        {'None': []},
        {
            'Đọc': ['bai_viet (Article)'],
            'Tạo': [],
            'Cập nhật': ['bai_viet.luot_xem (increment)'],
            'Xóa': []
        }
    )

    # 5. Xem và tra cứu địa chỉ quán
    add_listed_section(
        doc,
        '5. Xem và tra cứu địa chỉ quán',
        'Cho phép khách vãng lai xem danh sách các chi nhánh đang hoạt động và tra cứu chi tiết từng chi nhánh gồm địa chỉ, giờ mở/đóng, số điện thoại và đường dẫn bản đồ.',
        ['Hệ thống đang hoạt động.', 'Dữ liệu chi nhánh đã có trong CSDL, ít nhất một chi nhánh đang mở.'],
        ['Danh sách chi nhánh hoặc chi tiết được hiển thị; không có ghi thay đổi vào CSDL.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Truy cập mục Cửa hàng/Địa chỉ quán.', 'Lấy danh sách chi nhánh đang mở từ CSDL.'),
            ('Kiểm tra danh sách có dữ liệu không.', ''),
            ('Hiển thị danh sách: tên chi nhánh, địa chỉ rút gọn, giờ mở-đóng, số điện thoại.', ''),
            ('Chọn chi nhánh để xem chi tiết.', 'Kiểm tra chi nhánh tồn tại; trả về chi tiết; nếu có đường dẫn bản đồ, hiển thị nút Xem bản đồ.'),
            ('Nhấn Xem bản đồ để mở bản đồ.', 'Mở liên kết bản đồ hoặc đề xuất sao chép địa chỉ nếu không có link.')
        ],
        {
            '3.1 Không có chi nhánh đang hoạt động': ['Hiển thị thông báo tương ứng', 'Kết thúc use case.'],
            '6.1 Từ khóa tìm kiếm rỗng': ['Hiển thị toàn bộ danh sách', 'Quay lại bước 4.']
        },
        {
            '2.1 Lỗi CSDL khi tải danh sách': ['Hiển thị lỗi, ghi log, kết thúc.'],
            '14.1 Không thể mở ứng dụng bản đồ': ['Hiển thị thông báo; cho phép copy địa chỉ; ghi log.']
        },
        {'None': []},
        {
            'Đọc': ['chi_nhanh (Branch)'],
            'Tạo': [],
            'Cập nhật': [],
            'Xóa': []
        }
    )

    # 6. Nhắn tin hỗ trợ (Chatbox / AI / Zalo / Hotline)
    add_listed_section(
        doc,
        '6. Nhắn tin hỗ trợ (Chatbox / AI / Zalo / Hotline)',
        'Cho phép khách vãng lai liên hệ hỗ trợ qua 4 kênh: chatbox, trợ lý AI tích hợp, Zalo OA và gọi điện hotline. Cuộc hội thoại qua chatbox được lưu vào CSDL và thông báo đến nhân viên.',
        ['Hệ thống đang hoạt động.', 'Các kênh liên lạc đang trong trạng thái hoạt động.'],
        ['Chatbox: cuộc hội thoại và tin nhắn được lưu; AI: phản hồi hiển thị; Zalo/Hotline: ứng dụng/điện thoại được mở.'],
        'Khách vãng lai',
        'Không có',
        [
            ('Chọn biểu tượng Hỗ trợ.', 'Hiển thị giao diện lựa chọn kênh liên lạc.'),
            ('Chọn kênh mong muốn (Chatbox/AI/Zalo/Gọi).', 'Kiểm tra kênh hoạt động.'),
            # Chatbox subflow
            ('(Chatbox) Nhập nội dung tin nhắn và gửi.', 'Tạo `cuoc_hoi_thoai` (nếu mới) với trạng thái Đang mở; lưu `tin_nhan_hoi_thoai` vào CSDL; gửi thông báo tới nhân viên; hiển thị xác nhận gửi.'),
            # AI subflow
            ('(AI) Nhập câu hỏi và gửi.', 'Gửi câu hỏi đến dịch vụ AI; nhận và hiển thị phản hồi; (tùy cấu hình) ghi `nhat_ky_suy_luan`.')
        ],
        {
            '4.1 Chatbox ngoài giờ hoặc không có nhân viên': ['Hiển thị thông báo ngoài giờ; cho phép gửi tin; quay lại bước 5a.'],
            '6a.1 Tin nhắn rỗng (Chatbox)': ['Vô hiệu hóa nút gửi; quay lại bước 5a.'],
            '6b.1 Câu hỏi rỗng (AI)': ['Vô hiệu hóa nút gửi; quay lại bước 5b.']
        },
        {
            '7a.1 Lỗi CSDL khi tạo cuộc hội thoại': ['Hiển thị lỗi, không lưu; ghi log; kết thúc.'],
            '8a.1 Lỗi CSDL khi lưu tin nhắn': ['Hiển thị lỗi; ghi log; kết thúc.'],
            '7b.1 Lỗi kết nối dịch vụ AI': ['Hiển thị thông báo AI không khả dụng; ghi log; kết thúc.']
        },
        {
            'Subflow: Chat với nhân viên (Chatbox)': [
                'Tạo cuộc hội thoại mới (nếu cần).',
                'Lưu tin nhắn đầu tiên.',
                'Cập nhật trường tin nhắn cuối và thời gian của cuộc hội thoại.',
                'Gửi thông báo đến nhân viên phụ trách.'
            ],
            'Subflow: Chat với AI': [
                'Gửi câu hỏi tới endpoint AI.',
                'Nhận phản hồi; nếu AI không hiểu, đề xuất chuyển sang nhân viên.'
            ],
            'Subflow: Zalo / Gọi': [
                'Mở Zalo OA hoặc màn hình quay số tương ứng.'
            ]
        },
        {
            'Đọc': ['mo_hinh_ai (nếu cần để chọn mô hình)', 'cuoc_hoi_thoai (hiển thị danh sách cuộc hội thoại nếu có)'],
            'Tạo': ['cuoc_hoi_thoai', 'tin_nhan_hoi_thoai', 'nhat_ky_suy_luan (tùy cấu hình AI)'],
            'Cập nhật': ['cuoc_hoi_thoai.last_message, cuoc_hoi_thoai.updated_at'],
            'Xóa': []
        }
    )

    # Save file
    out = DOC_PATH
    doc.save(out)
    print(f"✅ File Word tạo thành công: {out}")

if __name__ == '__main__':
    main()
